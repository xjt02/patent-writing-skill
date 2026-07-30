"""
md_to_docx.py
=============

基于 assets/templates/专利空白模版_fixed.docx 的 5 节骨架，把 prompts/06_assemble.md
的所有上游产物（intake / abstract / IR / claims / description / formulas）注入对应节，
输出最终专利 .docx。

用法：
    python tools/md_to_docx.py \
      --template assets/templates/专利空白模版_fixed.docx \
      --case-id <YYYYMMDDHHmmss> \
      --abstract-page <abstract_page.md 或 ->段落文本> \
      --flowchart-png <path/to/fig.png> \
      --claims <claims.json> \
      --description <description.md> \
      --formulas <formulas.json> \
      --output assets/cases/<case_id>/<case_id>.docx

实现要点：
1. Document(template_path) — 不重置 docDefaults、sections、headers/footers
2. 节 1：在 sectPr 之前清空原模板的占位段，注入扉页摘要段落
3. 节 2：在节 1 sectPr 之后、节 2 sectPr 之前嵌入 PNG（居中）
4. 节 3：在节 2 sectPr 之后、节 3 sectPr 之前注入权利要求编号段
5. 节 4：在节 3 sectPr 之后、节 4 sectPr 之前渲染 description.md（替换 FORMULA 占位为 OMML）
6. 节 5：替换为「图 N  待贴附：XXX」占位段
7. 写文档属性（title / author / created）
"""
from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from lxml import etree


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"

# 识别 FORMULA 块中的 MathML 行（mrow/mi/mn/mo/msub/...，含闭合标签 </...>），
# 用于 md_to_docx.py 在渲染时跳过这些过程性产物，不作为正文段输出
_MATHML_TAG_RE = re.compile(
    r"^\s*</?(mrow|mi|mn|mo|msub|msup|msqrt|mfrac|munderover|mover|munder|"
    r"mtable|mtr|mtd|menclose|mfenced|mpadded|mstyle|ms|mspace)\b"
)

# 6 个 sectPr 把文档切成 5 节。模板正文段里第 1、第 2 段… 就是节切分处。
# 我们用「找 sectPr 节点」定位，不依赖段序号（模板可能微调）。


def _section_index(doc) -> int:
    """下一个可用的节计数。"""
    return len(doc.sections)


def _find_section_boundaries(doc) -> list[etree._Element]:
    """返回所有 sectPr 元素（含每节节内尾部 sectPr + 文档末尾 sectPr）。"""
    return doc.element.body.findall(f".//{W}sectPr")


def _section_last_paragraph_index(doc, sect_idx: int) -> int:
    """返回第 sect_idx 节末尾 sectPr 所在段（即节最后一段）的 body 子节点索引。"""
    body = doc.element.body
    sect_prs = body.findall(f".//{W}sectPr")
    if sect_idx >= len(sect_prs):
        raise IndexError(f"节 {sect_idx} 超出范围（仅 {len(sect_prs)} 节）")
    target_sect_pr = sect_prs[sect_idx]
    for i, child in enumerate(body):
        if child.tag != f"{W}p":
            continue
        if child.find(f"{W}pPr/{W}sectPr") is target_sect_pr:
            return i
    raise ValueError(f"找不到节 {sect_idx} 的 sectPr 所在段")


def _remove_paragraphs_in_section(doc, start_p_idx: int, end_p_idx: int) -> int:
    """
    删除 body[start_p_idx..end_p_idx) 中的段落（保留 end_p_idx 处的 sectPr 段）。
    返回删除的段数。
    """
    body = doc.element.body
    removed = 0
    children = list(body)
    # 只删段，不删 sectPr 段（end_p_idx）
    end_sect_pr_p = children[end_p_idx]
    for i in range(start_p_idx, end_p_idx):
        c = children[i]
        if c.tag == f"{W}p":
            body.remove(c)
            removed += 1
    # 重新数：被删后，end_p_idx 仍指向原 sectPr 段
    return removed


def _make_paragraph_after(doc, anchor_p: etree._Element, text: str = "") -> etree._Element:
    """
    在 anchor_p 之后插入一个新段落（anchor_p 通常是 sectPr 段）。
    返回新段元素。
    """
    new_p = etree.SubElement(anchor_p.getparent(), f"{W}p")
    anchor_p.addnext(new_p)
    if text:
        run = etree.SubElement(new_p, f"{W}r")
        rpr = etree.SubElement(run, f"{W}rPr")
        # 字体继承 docDefaults（不显式设置）
        t = etree.SubElement(run, f"{W}t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = text
    return new_p


def _set_first_line_indent(p: etree._Element, chars: int = 200) -> None:
    """设首行缩进 chars 个字符（200 = 2 字符）。"""
    ppr = p.find(f"{W}pPr")
    if ppr is None:
        ppr = etree.SubElement(p, f"{W}pPr")
        p.insert(0, ppr)
    ind = ppr.find(f"{W}ind")
    if ind is None:
        ind = etree.SubElement(ppr, f"{W}ind")
    ind.set(f"{W}firstLineChars", str(chars))
    ind.set(f"{W}firstLine", str(chars))


def _inject_section1(doc, abstract_text: str) -> None:
    """节 1 = 扉页摘要。

    模板里节 1 内是空白段。我们把那段删掉，注入 1 段首行缩进的摘要正文。
    """
    body = doc.element.body
    sect_prs = body.findall(f".//{W}sectPr")
    if not sect_prs:
        raise ValueError("模板无 sectPr")
    target = sect_prs[0]  # 第 1 节的 sectPr
    # 找包含该 sectPr 的段
    for child in body:
        if child.tag != f"{W}p":
            continue
        if child.find(f"{W}pPr/{W}sectPr") is target:
            sect1_end_p = child
            break
    else:
        raise ValueError("找不到节 1 sectPr 段")

    # 删 sect1_end_p 之前的所有段落（清空占位段）
    to_remove = []
    for child in body:
        if child is sect1_end_p:
            break
        if child.tag == f"{W}p":
            to_remove.append(child)
    for p in to_remove:
        body.remove(p)

    # 在 sect1_end_p 之前插入新段
    new_p = etree.Element(f"{W}p")
    sect1_end_p.addprevious(new_p)
    run = etree.SubElement(new_p, f"{W}r")
    t = etree.SubElement(run, f"{W}t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = abstract_text
    _set_first_line_indent(new_p, 200)


def _inject_section2(doc, png_path: Path) -> None:
    """节 2 = 摘要附图。删除占位段，嵌入 PNG（居中）。"""
    body = doc.element.body
    sect_prs = body.findall(f".//{W}sectPr")
    if len(sect_prs) < 2:
        raise ValueError("模板不足 2 节")

    sect1_end_p = _find_sect_end(body, 0)
    sect2_end_p = _find_sect_end(body, 1)
    _clear_between(sect1_end_p, sect2_end_p)

    # 先用 python-docx 在文档末尾添加带图段（自动处理 rId / relationships）
    # 再把那段 move 到 sect2_end_p 之前
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.add_run().add_picture(str(png_path), width=Inches(5.0))
    # 取出刚添加的段（最后一个 body 子节点），move 到 sect2_end_p 之前
    img_p_elem = img_para._p
    body.remove(img_p_elem)
    sect2_end_p.addprevious(img_p_elem)


def _find_sect_end(body, idx: int):
    """第 idx 节的「末尾标记」：sectPr 在段里就返回段，sectPr 直接挂 body 就返回 sectPr 本身。

    模板里节 5 的 sectPr 直接挂在 body 末尾（无占位段）。
    """
    sect_prs = body.findall(f".//{W}sectPr")
    target = sect_prs[idx]
    parent = target.getparent()
    if parent is body:
        return target
    for child in body:
        if child.tag == f"{W}p" and child.find(f"{W}pPr/{W}sectPr") is target:
            return child
    raise ValueError(f"找不到节 {idx} 的末尾段")


def _inject_section3(doc, claims: list[dict]) -> None:
    """节 3 = 权利要求书。删除占位段，注入编号段。

    独立权利要求（id=1）：拆段 = 前导段（不缩进）+ 步骤段（每段缩进 2 字符）
    从属权利要求（id=2..5）：保持单段，不缩进
    """
    body = doc.element.body
    sect2_end_p = _find_sect_end(body, 1)
    sect3_end_p = _find_sect_end(body, 2)
    _clear_between(sect2_end_p, sect3_end_p)

    # 独立权利要求步骤拆分正则：按「；」+ 步骤 N 拆分（lookahead 不消耗分隔符）
    _STEP_SPLIT_RE = re.compile(r"；(?=步骤\s*\d+、)")

    for i, claim in enumerate(claims, 1):
        text = claim.get("text", "").strip()
        if not text:
            continue

        # 独立权利要求：拆段（前导不缩进 + 步骤每段缩进 2 字符）
        if i == 1 and "步骤 1、" in text:
            anchor = "包括："
            anchor_idx = text.find(anchor)
            if anchor_idx != -1:
                prefix = text[: anchor_idx + len(anchor)]  # "一种XXX方法，其特征在于，包括："
                steps_str = text[anchor_idx + len(anchor) :]  # "步骤 1、A；步骤 2、B；...；步骤 5、E。"
                steps = _STEP_SPLIT_RE.split(steps_str)
                # 第 1 段：1、 + 前缀（顶格，不缩进）
                new_p = etree.Element(f"{W}p")
                sect3_end_p.addprevious(new_p)
                run = etree.SubElement(new_p, f"{W}r")
                t = etree.SubElement(run, f"{W}t")
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                t.text = f"1、{prefix}"
                # 后续每段：每个步骤独立成段，首行缩进 2 字符
                for step in steps:
                    step = step.strip()
                    if not step:
                        continue
                    new_p = etree.Element(f"{W}p")
                    sect3_end_p.addprevious(new_p)
                    run = etree.SubElement(new_p, f"{W}r")
                    t = etree.SubElement(run, f"{W}t")
                    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                    t.text = step
                    _set_first_line_indent(new_p, 200)  # 段首缩进 2 字符
                continue
            # 找不到「包括：」时 fallback 到原逻辑

        # 从属权利要求（id=2..5）或独立权利要求 fallback：单段，不缩进
        new_p = etree.Element(f"{W}p")
        sect3_end_p.addprevious(new_p)
        run = etree.SubElement(new_p, f"{W}r")
        # "1、 " 前缀 + 权利要求正文
        full = f"{i}、{text}"
        t = etree.SubElement(run, f"{W}t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = full
        # 段首不缩进（从属权利要求）；行距 line=360 lineRule=auto 继承自 docDefaults


def _clear_between(start_p, end_p) -> None:
    """清空 start_p 与 end_p 之间的段落（不含两者本身）。"""
    parent = start_p.getparent()
    to_remove = []
    node = start_p.getnext()
    while node is not None and node is not end_p:
        if node.tag == f"{W}p":
            to_remove.append(node)
        node = node.getnext()
    for n in to_remove:
        parent.remove(n)


def _inject_section4(doc, description_md: str, formulas: dict[str, str]) -> None:
    """节 4 = 说明书。渲染 description.md 段，替换 <<<FORMULA:n>>> 为 OMML。

    description.md 格式：5 小节标题以 `## 技术领域` `## 背景技术` 等开头。
    公式占位符：<<FORMULA:NAME>> （单括号以匹配 04_description.md 提到的占位惯例）
    """
    body = doc.element.body
    sect3_end_p = _find_sect_end(body, 2)
    sect4_end_p = _find_sect_end(body, 3)
    _clear_between(sect3_end_p, sect4_end_p)

    # 锚点固定为 sect4_end_p（节 4 末尾段）。
    # addprevious 语义：把新段插到 anchor 紧邻前面；连续调用产生正序累积。
    # 若循环里更新 current_p，新段会插到上一次新段之前，导致内容逆序。
    anchor = sect4_end_p
    # 国标：每个小节内容结束后空一行。在下一个 ## 标题处理前 / 循环结束时插入空段
    last_section_had_text = False
    # 用 sys.path 把同目录的 mathml_to_omml 拉进来
    this_dir = Path(__file__).parent
    if str(this_dir) not in sys.path:
        sys.path.insert(0, str(this_dir))
    from mathml_to_omml import mathml_to_omml, inject_into_paragraph
    from docx.text.paragraph import Paragraph

    def _new_p_after(prev_p, text: str = "") -> etree._Element:
        new_p = etree.Element(f"{W}p")
        prev_p.addprevious(new_p)
        if text:
            run = etree.SubElement(new_p, f"{W}r")
            t = etree.SubElement(run, f"{W}t")
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = text
        return new_p

    lines = description_md.splitlines()
    for line in lines:
        s = line.rstrip()
        if not s:
            continue
        # 说明书节标题：# 一种XXX的YYY方法
        if s.startswith("# "):
            new_p = _new_p_after(anchor)
            run = etree.SubElement(new_p, f"{W}r")
            rpr = etree.SubElement(run, f"{W}rPr")
            b = etree.SubElement(rpr, f"{W}b")  # 国标：说明书标题加粗
            t = etree.SubElement(run, f"{W}t")
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = s[2:].strip()
            # 段后 1 行间距 ≈ 18pt = 360 twips
            ppr = etree.SubElement(new_p, f"{W}pPr")
            new_p.insert(0, ppr)
            spacing = etree.SubElement(ppr, f"{W}spacing")
            spacing.set(f"{W}after", "360")
            continue
        # 小节标题：## 名
        if s.startswith("## "):
            # 上一小节段尾空一行（如果不是第一个小节）
            if last_section_had_text:
                _new_p_after(anchor)  # 空段
            last_section_had_text = False
            new_p = _new_p_after(anchor)
            run = etree.SubElement(new_p, f"{W}r")
            rpr = etree.SubElement(run, f"{W}rPr")
            b = etree.SubElement(rpr, f"{W}b")
            # 国标 GB/T 专利说明书：5 小节标题必须加粗 + 加下划线 + 顶格
            u = etree.SubElement(rpr, f"{W}u")
            u.set(f"{W}val", "single")
            t = etree.SubElement(run, f"{W}t")
            t.text = s[3:].strip()
            continue
        # FORMULA 占位
        m = re.match(r"^<<FORMULA:(\w+)>>$", s)
        if m:
            key = m.group(1)
            if key in formulas:
                # 国标：公式前后不要空一行，仅插入公式段
                new_p = _new_p_after(anchor)
                para = Paragraph(new_p, None)
                inject_into_paragraph(para, formulas[key])
            last_section_had_text = True  # 公式也算小节内容
            continue
        # 过滤 FORMULA 块过程性产物（MathML 行 / <<END>> / </math>）
        # 这些行只服务于 tools/extract_formulas.py 抽取公式，
        # 在 docx 中只渲染为可编辑 OMML 公式段，不展示原始标记
        stripped = s.lstrip()
        if (stripped.startswith("<math") or stripped.startswith("</math>")
                or stripped.startswith("<<END") or _MATHML_TAG_RE.match(s)):
            continue
        # 普通正文段
        new_p = _new_p_after(anchor, s)
        _set_first_line_indent(new_p, 200)
        last_section_had_text = True

    # 循环结束后，最后一个小节段尾也空一行
    if last_section_had_text:
        _new_p_after(anchor)


def _inject_section5(doc, placeholders: list[str]) -> None:
    """节 5 = 说明书附图。仅占位段。"""
    body = doc.element.body
    sect4_end_p = _find_sect_end(body, 3)
    sect5_end_p = _find_sect_end(body, 4)
    _clear_between(sect4_end_p, sect5_end_p)
    current_p = sect4_end_p
    for i, desc in enumerate(placeholders, 2):  # 图 2 起
        new_p = etree.Element(f"{W}p")
        sect5_end_p.addprevious(new_p)
        run = etree.SubElement(new_p, f"{W}r")
        t = etree.SubElement(run, f"{W}t")
        t.text = f"图{i}  待贴附：{desc}"
        current_p = new_p


def _set_doc_properties(doc, title: str, author: str) -> None:
    cp = doc.core_properties
    cp.title = title
    cp.author = author


def main(argv=None):
    parser = argparse.ArgumentParser(description="5 节结构组装：注入摘要/图/权利要求/说明书/附图")
    parser.add_argument("--template", type=Path, required=True)
    parser.add_argument("--case-id", type=str, required=True)
    parser.add_argument("--abstract-page", type=Path, help="扉页摘要 .md（首行为正文）")
    parser.add_argument("--abstract-text", type=str, help="或直接传入摘要文本")
    parser.add_argument("--flowchart-png", type=Path, help="节 2 嵌入 PNG")
    parser.add_argument("--claims", type=Path, help="claims.json 路径")
    parser.add_argument("--description", type=Path, help="说明书 .md 路径")
    parser.add_argument("--formulas", type=Path, help="formulas.json 路径")
    parser.add_argument("--placeholders", type=str, nargs="*",
                        help="节 5 占位段描述（默认 '附图'）")
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--doc-title", type=str, default="")
    parser.add_argument("--doc-author", type=str, default="")
    args = parser.parse_args(argv)

    if not args.template.is_file():
        sys.exit(f"模板不存在: {args.template}")
    args.output.parent.mkdir(parents=True, exist_ok=True)

    # 复制模板到 output
    shutil.copy(args.template, args.output)
    doc = Document(str(args.output))

    # 收集输入
    abstract = args.abstract_text or ""
    if args.abstract_page and args.abstract_page.is_file():
        abstract = args.abstract_page.read_text(encoding="utf-8").strip()

    png_path = args.flowchart_png
    claims: list[dict] = []
    if args.claims and args.claims.is_file():
        claims = json.loads(args.claims.read_text(encoding="utf-8"))
        # 兼容 {"claims":[...]} 与裸 [...] 两种 schema
        if isinstance(claims, dict) and "claims" in claims:
            claims = claims["claims"]

    description = ""
    if args.description and args.description.is_file():
        description = args.description.read_text(encoding="utf-8")

    formulas: dict[str, str] = {}
    if args.formulas and args.formulas.is_file():
        formulas = json.loads(args.formulas.read_text(encoding="utf-8"))

    placeholders = args.placeholders or ["附图"]

    # 注入（顺序至关重要：不能反，先动节 1，最后是节 5）
    if abstract:
        _inject_section1(doc, abstract)
    if png_path and png_path.is_file():
        _inject_section2(doc, png_path)
    if claims:
        _inject_section3(doc, claims)
    if description:
        _inject_section4(doc, description, formulas)
    _inject_section5(doc, placeholders)

    if args.doc_title or args.doc_author:
        _set_doc_properties(doc, args.doc_title, args.doc_author)

    doc.save(str(args.output))
    print(f"已写入: {args.output}")
    print(f"  节 1 摘要: {len(abstract)} 字")
    print(f"  节 2 PNG: {'已嵌入' if png_path else '无'}")
    print(f"  节 3 权利要求: {len(claims)} 条")
    print(f"  节 4 说明书: {description.count(chr(10))} 行, {len(formulas)} 公式")
    print(f"  节 5 占位段: {len(placeholders)} 个")


if __name__ == "__main__":
    main()