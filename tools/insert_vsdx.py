"""
insert_vsdx.py
==============

把 flowchart IR（visio-scientific-flowchart 格式）转 vsdx + PNG，
PNG 嵌入 docx 节 2，vsdx 作为附带源文件交付。

用法：
    python tools/insert_vsdx.py <flowchart_ir.json> <template.docx> <output_dir> --case-id ID

输出：
    - <output_dir>/<case_id>.vsdx       流程图源文件
    - <output_dir>/<case_id>.png        PNG 嵌入图
    - <output_dir>/<case_id>.docx       含嵌入 PNG 的 docx

依赖：
- visio-scientific-flowchart（已装）
- pywin32（Visio COM 调用）
- Microsoft Visio（Windows + 已安装）
- python-docx + lxml

PNG 生成走 Visio COM 路线：先用 visio-scientific-flowchart 出 vsdx，
再用 Visio COM 打开该 vsdx → ExportAsFixedFormat(2, ...) 出 PNG。
这样无需安装 cairosvg 或其他 SVG 后端。
"""
from __future__ import annotations

import argparse
import json
import os
import shutil
import sys
from pathlib import Path

# visio-scientific-flowchart 路径：优先用环境变量，否则相对于本脚本位置向上查找
# 用法：设置 VISIO_SKILL_PATH 环境变量指向 visio-scientific-flowchart 目录
_env_path = os.environ.get("VISIO_SKILL_PATH")
if _env_path:
    VSDX_SKILL_TPL = Path(_env_path) / "templates"
else:
    # 相对于本脚本位置向上两级（tools/ → skill 根）
    VSDX_SKILL_TPL = Path(__file__).parent.parent / "visio-scientific-flowchart" / "templates"

sys.path.insert(0, str(VSDX_SKILL_TPL))

try:
    from ir_adapter import to_vsdx
except ImportError as e:
    print(f"无法导入 visio-scientific-flowchart: {e}", file=sys.stderr)
    raise

try:
    import win32com.client
    import pythoncom
    HAS_COM = True
except ImportError:
    HAS_COM = False
    print("缺少 pywin32，请运行: pip install pywin32", file=sys.stderr)

from docx import Document
from docx.shared import Inches


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"


def generate_vsdx(ir: dict, output_path: Path) -> bool:
    """调 visio-scientific-flowchart 出 vsdx。返回是否成功。"""
    try:
        spec = to_vsdx(ir, str(output_path))
        spec.save()
        return output_path.exists() and output_path.stat().st_size > 0
    except Exception as e:
        print(f"  vsdx 生成失败: {type(e).__name__}: {e}", file=sys.stderr)
        return False


def export_png_via_visio_com(vsdx_path: Path, png_path: Path) -> bool:
    """
    用 Visio COM 重新打开 vsdx 并导出 PNG。

    实现：
        Documents.Open(vsdx) → Pages.Item(1) → Page.Export(png)
        Visio 按页面尺寸（默认 1080x613）导出 PNG。

    已知陷阱：
        - Pages 是动态集合，__getitem__ 会失败；必须用 Pages.Item(idx)
        - Documents.Open 可能因 Visio 进程残留文件锁而失败
          → 调用方应先 import cleanup_visio() 杀掉残留 Visio.exe
        - Page.Export(FileName) 第二个参数是关于是否覆盖（可选）
    """
    if not HAS_COM:
        return False
    try:
        pythoncom.CoInitialize()
        visio = win32com.client.Dispatch("Visio.Application")
        visio.Visible = False
        visio.AlertResponse = 7  # 禁用 Visio 弹窗
        try:
            doc = visio.Documents.Open(str(vsdx_path))
            try:
                page = doc.Pages.Item(1)
                page.Export(str(png_path))
            finally:
                doc.Close()
        finally:
            visio.Quit()
            pythoncom.CoUninitialize()
        return png_path.exists() and png_path.stat().st_size > 0
    except Exception as e:
        print(f"  Visio COM 导出 PNG 失败: {type(e).__name__}: {e}", file=sys.stderr)
        return False


def cleanup_visio() -> None:
    """
    杀掉残留的 VISIO.EXE 进程（避免 Documents.Open 文件共享冲突）。

    用 subprocess 而非 bash 调 taskkill，避开 Git Bash 把 '/F/' 当路径。
    """
    if not HAS_COM:
        return
    try:
        import subprocess
        subprocess.run(
            ["taskkill", "/F", "/IM", "VISIO.EXE"],
            capture_output=True,
            text=True,
        )
    except Exception as e:
        print(f"  清理 Visio 进程失败: {e}", file=sys.stderr)


def generate_png_from_vsdx(vsdx_path: Path, png_path: Path) -> bool:
    """从已生成的 vsdx 文件导出 PNG（Visio COM 路线）。

    步骤：
        1. 清理残留 Visio.exe 进程（避免文件锁冲突）
        2. Documents.Open(vsdx)
        3. Pages.Item(1).Export(png)
        4. 关闭 Visio 实例
    """
    cleanup_visio()
    return export_png_via_visio_com(vsdx_path, png_path)


def find_section1_last_paragraph(doc):
    """
    找到节 1（扉页摘要）的最后一段（含 sectPr 标记节 1 结束）。
    节 2 的所有段落都在这一段之后。
    """
    body = doc.element.body
    for child in body:
        if child.tag == f"{W}p":
            pPr = child.find(f"{W}pPr")
            if pPr is not None and pPr.find(f"{W}sectPr") is not None:
                from docx.text.paragraph import Paragraph
                return Paragraph(child, doc.paragraphs[0]._parent)
    raise ValueError("找不到节 1 的 sectPr，无法定位节 2 插入点")


def insert_png_into_section2(doc, png_path: Path, anchor_para) -> None:
    """
    把 PNG 插入节 2（在 anchor_para 之后插入新段含图片）。

    实现：用 add_picture 加到末尾，再把含图片的段移动到 anchor 之后。
    """
    doc.add_picture(str(png_path), width=Inches(5.0))
    last_para = doc.paragraphs[-1]
    last_p_elem = last_para._p
    # 从原位置删除
    last_p_elem.getparent().remove(last_p_elem)
    # 插入到 anchor_para 之后（即节 2 内）
    anchor_para._p.addnext(last_p_elem)


def main(argv=None):
    parser = argparse.ArgumentParser(description="生成 vsdx + PNG 并嵌入 docx 节 2")
    parser.add_argument("ir_json", type=Path, help="flowchart IR JSON 路径")
    parser.add_argument("template", type=Path, help="模板 docx 路径")
    parser.add_argument("output_dir", type=Path, help="输出目录")
    parser.add_argument("--case-id", type=str, required=True, help="案件 ID")
    parser.add_argument("--no-vsdx", action="store_true", help="跳过 vsdx 生成")
    args = parser.parse_args(argv)

    if not args.ir_json.is_file():
        sys.exit(f"IR JSON 不存在: {args.ir_json}")
    if not args.template.is_file():
        sys.exit(f"模板不存在: {args.template}")

    args.output_dir.mkdir(parents=True, exist_ok=True)
    ir = json.loads(args.ir_json.read_text(encoding="utf-8"))
    # 国标：摘要附图不绘制流程图标题，标题由附图说明节单独列出
    if isinstance(ir.get("meta"), dict):
        ir["meta"]["title"] = ""

    vsdx_path = args.output_dir / f"{args.case_id}.vsdx"
    png_path = args.output_dir / f"{args.case_id}.png"
    docx_path = args.output_dir / f"{args.case_id}.docx"

    # 1. 生成 vsdx
    if not args.no_vsdx:
        print(f"[1/3] 生成 vsdx: {vsdx_path}")
        if not generate_vsdx(ir, vsdx_path):
            print("  ⚠ vsdx 生成失败，继续生成 PNG")
    else:
        print("[1/3] 跳过 vsdx 生成（--no-vsdx）")

    # 2. 生成 PNG（从 vsdx 用 Visio COM 导出）
    print(f"[2/3] 生成 PNG: {png_path}")
    if not vsdx_path.exists():
        sys.exit("vsdx 不存在，无法导出 PNG")
    if not generate_png_from_vsdx(vsdx_path, png_path):
        sys.exit("PNG 生成失败（Visio COM 路径），无法继续")

    # 3. 复制模板 → 插入 PNG → 保存
    print(f"[3/3] 组装 docx: {docx_path}")
    shutil.copy(args.template, docx_path)
    doc = Document(str(docx_path))
    anchor_para = find_section1_last_paragraph(doc)
    insert_png_into_section2(doc, png_path, anchor_para)
    doc.save(str(docx_path))

    print(f"\n✅ 完成: {docx_path}")
    print(f"  vsdx: {vsdx_path} ({vsdx_path.stat().st_size if vsdx_path.exists() else 0} bytes)")
    print(f"  PNG:  {png_path} ({png_path.stat().st_size} bytes)")
    print(f"  docx: {docx_path}")


if __name__ == "__main__":
    main()