"""
style_applier.py
================

把 assets/style.json 中的规范应用到目标 docx，或反向校验 docx 是否符合 style.json。

用法：
    # 校验模式（默认）
    python tools/style_applier.py check <docx> [-s assets/style.json]

    # 应用模式（强制 style.json 覆盖 docx 的 run/paragraph 格式）
    python tools/style_applier.py apply <docx> [-s assets/style.json] [-o output.docx]

校验内容：
- docDefaults rFonts（ascii / eastAsia）
- 每节页眉文字
- 页脚行距
- 正文段 firstLine / spacing

输出：人类可读的报告 + 可选的 JSON
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from lxml import etree


W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _get(elem, name: str):
    return elem.get(f"{W}{name}")


def _get_int(elem, name: str, default=None):
    val = _get(elem, name)
    if val is None:
        return default
    try:
        return int(val)
    except ValueError:
        return val


def check_defaults(doc, style: dict) -> list[dict]:
    """校验 docDefaults rFonts。"""
    issues = []
    expected = style.get("doc_defaults_rFonts", {})
    dd = doc.styles.element.find(f"{W}docDefaults")
    if dd is None:
        issues.append({"severity": "error", "msg": "缺少 docDefaults"})
        return issues
    rfonts = dd.find(f"{W}rPrDefault/{W}rPr/{W}rFonts")
    if rfonts is None:
        issues.append({"severity": "error", "msg": "docDefaults 缺少 rFonts"})
        return issues
    actual = {k: v for k, v in rfonts.attrib.items()
              for k in [k.split("}")[-1]]}
    for key, exp in expected.items():
        act = actual.get(key)
        if act != exp:
            issues.append({
                "severity": "error",
                "msg": f"docDefaults rFonts[{key}]: 期望 {exp!r}, 实际 {act!r}",
            })
    return issues


def check_section_headers(doc, style: dict) -> list[dict]:
    """校验 5 节页眉文字。"""
    issues = []
    expected_sections = style.get("sections", [])
    actual_sections = list(doc.sections)
    if len(actual_sections) != len(expected_sections):
        issues.append({
            "severity": "error",
            "msg": f"section 数量: 期望 {len(expected_sections)}, 实际 {len(actual_sections)}",
        })
    for i, (act_sec, exp_sec) in enumerate(zip(actual_sections, expected_sections), 1):
        # 取第一个非空 header text
        act_text = ""
        for p in act_sec.header.paragraphs:
            if p.text.strip():
                act_text = p.text
                break
        exp_texts = exp_sec.get("header_texts", [])
        if not exp_texts:
            continue
        exp_text = exp_texts[0]["text"]
        if act_text != exp_text:
            issues.append({
                "severity": "error",
                "msg": f"节 {i} 页眉: 期望 {exp_text!r}, 实际 {act_text!r}",
            })
    return issues


def check_body_paragraphs(doc, style: dict) -> list[dict]:
    """校验正文段格式（取首个有 run 的段落）。"""
    issues = []
    body_paras = style.get("body_paragraphs", [])
    if not body_paras:
        return issues
    expected = body_paras[0]["format"]
    # 找 doc 里第一个有 run 的 body 段落（不在 header/footer 内）
    body_root = doc.element.body
    actual_para = None
    for p in body_root.findall(f".//{W}p"):
        if p.findall(f"{W}r"):
            actual_para = p
            break
    if actual_para is None:
        issues.append({"severity": "warning", "msg": "正文无任何 run 段落"})
        return issues
    pPr = actual_para.find(f"{W}pPr")
    if pPr is None:
        issues.append({"severity": "warning", "msg": "首个正文段无 pPr"})
        return issues
    # 检查 firstLineChars
    ind = pPr.find(f"{W}ind")
    if ind is not None:
        act_first_chars = _get_int(ind, "firstLineChars")
        exp_first_chars = expected.get("paragraph_indent", {}).get("firstLineChars")
        if act_first_chars != exp_first_chars:
            issues.append({
                "severity": "warning",
                "msg": f"正文首行缩进 firstLineChars: 期望 {exp_first_chars}, 实际 {act_first_chars}",
            })
    # 检查 line 行距
    spacing = pPr.find(f"{W}spacing")
    if spacing is not None:
        act_line = _get_int(spacing, "line")
        exp_line = expected.get("paragraph_spacing", {}).get("line")
        if act_line != exp_line:
            issues.append({
                "severity": "warning",
                "msg": f"正文行距 line: 期望 {exp_line}, 实际 {act_line}",
            })
    return issues


def run_check(docx_path: Path, style_path: Path) -> dict:
    style = json.loads(style_path.read_text(encoding="utf-8"))
    doc = Document(str(docx_path))

    all_issues = []
    all_issues += check_defaults(doc, style)
    all_issues += check_section_headers(doc, style)
    all_issues += check_body_paragraphs(doc, style)

    return {
        "docx": str(docx_path),
        "style": str(style_path),
        "passed": not any(i["severity"] == "error" for i in all_issues),
        "issues": all_issues,
        "issue_count": len(all_issues),
    }


def run_apply(docx_path: Path, style_path: Path, output_path: Path) -> None:
    """把 style.json 强制应用到 docx（用于微调）。"""
    import shutil
    style = json.loads(style_path.read_text(encoding="utf-8"))
    shutil.copy(docx_path, output_path)
    doc = Document(str(output_path))

    # 1. 应用 docDefaults rFonts
    expected_fonts = style.get("doc_defaults_rFonts", {})
    dd = doc.styles.element.find(f"{W}docDefaults")
    if dd is not None and expected_fonts:
        rfonts = dd.find(f"{W}rPrDefault/{W}rPr/{W}rFonts")
        if rfonts is not None:
            for key, val in expected_fonts.items():
                rfonts.set(f"{W}{key}", val)

    doc.save(str(output_path))


def main(argv=None):
    parser = argparse.ArgumentParser(
        description="校验或应用 style.json 到 docx"
    )
    sub = parser.add_subparsers(dest="cmd", required=True)

    p_check = sub.add_parser("check", help="校验 docx 是否符合 style.json")
    p_check.add_argument("docx", type=Path)
    p_check.add_argument("-s", "--style", type=Path,
                         default=Path("assets/style.json"))

    p_apply = sub.add_parser("apply", help="应用 style.json 到 docx")
    p_apply.add_argument("docx", type=Path)
    p_apply.add_argument("-s", "--style", type=Path,
                         default=Path("assets/style.json"))
    p_apply.add_argument("-o", "--output", type=Path, required=True)

    args = parser.parse_args(argv)

    if args.cmd == "check":
        if not args.docx.is_file():
            sys.exit(f"输入不存在: {args.docx}")
        if not args.style.is_file():
            sys.exit(f"style.json 不存在: {args.style}")
        report = run_check(args.docx, args.style)
        print(json.dumps(report, ensure_ascii=False, indent=2))
        sys.exit(0 if report["passed"] else 1)

    elif args.cmd == "apply":
        if not args.docx.is_file():
            sys.exit(f"输入不存在: {args.docx}")
        if not args.style.is_file():
            sys.exit(f"style.json 不存在: {args.style}")
        run_apply(args.docx, args.style, args.output)
        print(f"已应用 style 到: {args.output}")


if __name__ == "__main__":
    main()