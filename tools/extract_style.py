"""
extract_style.py
================

从专利空白模板 .docx 抽取排版规范，输出 style.json。

用法：
    python tools/extract_style.py <input.docx> [-o style.json]

抽取内容：
- 5 节顺序与每节页眉/页脚文字
- 字体配置（ascii / hAnsi / eastAsia / cs）
- 字号（w:sz 与 w:szCs，半磅为单位）
- 行距（w:spacing/@w:line + @w:lineRule）
- 首行缩进（w:ind/@w:firstLine + @w:firstLineChars）
- 对齐方式（w:jc）
- 页眉/页脚段落的 szCs / firstLine

输出：
- assets/style.json（默认）
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn


def _w(tag: str) -> str:
    """WordprocessingML 命名空间简写。"""
    return qn(f"w:{tag}")


def _get_int(elem, name: str, default=None):
    """读取 w:xxx 属性，优先转 int；非数字（如 lineRule="exact"）保留为 str。"""
    val = elem.get(_w(name))
    if val is None:
        return default
    try:
        return int(val)
    except ValueError:
        return val


def _get_text(elem) -> str:
    """读取 w:t 文本节点（合并所有 run）。"""
    texts = []
    for t in elem.iter(_w("t")):
        if t.text:
            texts.append(t.text)
    return "".join(texts)


def extract_doc_defaults(doc) -> dict[str, str]:
    """docDefaults 的 rFonts。"""
    dd = doc.styles.element.find(_w("docDefaults"))
    if dd is None:
        return {}
    rfonts = dd.find(f"{_w('rPrDefault')}/{_w('rPr')}/{_w('rFonts')}")
    if rfonts is None:
        return {}
    return {
        attr.split("}")[-1]: val
        for attr, val in rfonts.attrib.items()
        if val
    }


def extract_paragraph_format(p) -> dict[str, Any]:
    """抽取单个段落格式：rPr (font) + pPr (indent/line/jc) + runs 的 sz/szCs。"""
    pPr = p._p.find(_w("pPr"))
    result: dict[str, Any] = {}

    # 段落 rPr（段落默认 rPr）
    if pPr is not None:
        pPr_rPr = pPr.find(_w("rPr"))
        if pPr_rPr is not None:
            rfonts = pPr_rPr.find(_w("rFonts"))
            if rfonts is not None:
                result["paragraph_rFonts"] = {
                    k.split("}")[-1]: v for k, v in rfonts.attrib.items()
                }
            sz = _get_int(pPr_rPr, "sz")
            if sz is not None:
                result["paragraph_sz_half_pt"] = sz
            szCs = _get_int(pPr_rPr, "szCs")
            if szCs is not None:
                result["paragraph_szCs_half_pt"] = szCs

        # spacing
        spacing = pPr.find(_w("spacing"))
        if spacing is not None:
            sp = {}
            for attr in ("line", "lineRule", "before", "after", "beforeLines", "afterLines"):
                v = _get_int(spacing, attr)
                if v is not None:
                    sp[attr] = v
            if sp:
                result["paragraph_spacing"] = sp

        # indent
        ind = pPr.find(_w("ind"))
        if ind is not None:
            ind_d = {}
            for attr in ("firstLine", "firstLineChars", "left", "leftChars",
                         "right", "rightChars", "hanging", "hangingChars"):
                v = _get_int(ind, attr)
                if v is not None:
                    ind_d[attr] = v
            if ind_d:
                result["paragraph_indent"] = ind_d

        # jc
        jc = pPr.find(_w("jc"))
        if jc is not None:
            result["paragraph_jc"] = jc.get(_w("val"))

    # run 级 sz / szCs（取首个非空）
    for r in p._p.findall(_w("r")):
        rPr = r.find(_w("rPr"))
        if rPr is None:
            continue
        sz = _get_int(rPr, "sz")
        if sz and "first_run_sz_half_pt" not in result:
            result["first_run_sz_half_pt"] = sz
        szCs = _get_int(rPr, "szCs")
        if szCs and "first_run_szCs_half_pt" not in result:
            result["first_run_szCs_half_pt"] = szCs
        # rFonts hint
        rfonts = rPr.find(_w("rFonts"))
        if rfonts is not None:
            hint = rfonts.get(_w("hint"))
            if hint and "first_run_font_hint" not in result:
                result["first_run_font_hint"] = hint
        break

    return result


def extract_section(s, idx: int) -> dict[str, Any]:
    """抽取一个 section 的页眉/页脚文字与字号。"""
    sec = {"index": idx, "header_texts": [], "footer_texts": []}

    # 页眉（每段一行）
    for p in s.header.paragraphs:
        txt = p.text
        if txt.strip() or txt:  # 包括纯空格的页眉
            sz_info = extract_paragraph_format(p)
            sec["header_texts"].append({"text": txt, "format": sz_info})

    # 页脚（每段一行）
    for p in s.footer.paragraphs:
        txt = p.text
        if txt or any(t.text for t in p._p.iter(_w("t")) if t.text):
            sz_info = extract_paragraph_format(p)
            sec["footer_texts"].append({"text": txt, "format": sz_info})

    return sec


def extract_style(docx_path: Path) -> dict[str, Any]:
    doc = Document(str(docx_path))
    result: dict[str, Any] = {
        "source_file": str(docx_path),
        "section_count": len(doc.sections),
        "doc_defaults_rFonts": extract_doc_defaults(doc),
        "sections": [extract_section(s, i + 1) for i, s in enumerate(doc.sections)],
        "body_paragraphs": [],
    }

    # 正文段落格式（取有 run 的前 10 个）
    for i, p in enumerate(doc.paragraphs[:10]):
        pf = extract_paragraph_format(p)
        if pf:
            result["body_paragraphs"].append({"index": i, "format": pf})

    return result


def main(argv=None):
    parser = argparse.ArgumentParser(description="从专利模板 docx 抽取排版规范为 JSON")
    parser.add_argument("docx", type=Path, help="输入 docx 路径")
    parser.add_argument("-o", "--output", type=Path,
                        default=Path("assets/style.json"),
                        help="输出 JSON 路径（默认 assets/style.json）")
    args = parser.parse_args(argv)

    if not args.docx.is_file():
        sys.exit(f"输入文件不存在: {args.docx}")

    data = extract_style(args.docx)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"已抽取 {data['section_count']} 节，输出到 {args.output}")


if __name__ == "__main__":
    main()